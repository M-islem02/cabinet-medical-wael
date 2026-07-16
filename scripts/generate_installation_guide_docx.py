from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Guide_Installation_MedCareSO_Cabinet_A_Z.docx"
SKILL_SCRIPTS = Path(
    r"C:\Users\lenovo\.codex\plugins\cache\openai-primary-runtime\documents\26.514.12219\skills\documents\scripts"
)
sys.path.insert(0, str(SKILL_SCRIPTS))
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


NAVY = "183B56"
BLUE = "236A9B"
LIGHT_BLUE = "EAF3F8"
PALE_BLUE = "F5F9FC"
GREEN = "1F7A5A"
LIGHT_GREEN = "EAF6F0"
AMBER = "B66A12"
LIGHT_AMBER = "FFF4DE"
RED = "A73535"
LIGHT_RED = "FCECEC"
INK = "202C39"
MUTED = "5F6F7F"
LINE = "CBD9E4"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = LINE, size: int = 6) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char_1, instr_text, fld_char_2])


def add_numbering_definition(doc: Document, kind: str) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    lvl.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if kind == "bullet" else "%1.")
    lvl.append(lvl_text)
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    lvl.append(suff)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    p_pr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    lvl.append(p_pr)
    if kind == "bullet":
        r_pr = OxmlElement("w:rPr")
        fonts = OxmlElement("w:rFonts")
        fonts.set(qn("w:ascii"), "Arial")
        fonts.set(qn("w:hAnsi"), "Arial")
        r_pr.append(fonts)
        lvl.append(r_pr)
    abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abs_ref = OxmlElement("w:abstractNumId")
    abs_ref.set(qn("w:val"), str(abstract_id))
    num.append(abs_ref)
    numbering.append(num)
    return num_id


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_list(doc: Document, items: list[str], num_id: int, numbered: bool = False) -> None:
    for item in items:
        p = doc.add_paragraph(style="Body Compact")
        apply_num(p, num_id)
        p.add_run(item)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(text, style=f"Heading {level}")
    keep_with_next(p)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Code Block")
    p.add_run(text)


def add_callout(doc: Document, title: str, body: str, tone: str = "info") -> None:
    fills = {"info": LIGHT_BLUE, "ok": LIGHT_GREEN, "warn": LIGHT_AMBER, "danger": LIGHT_RED}
    accents = {"info": BLUE, "ok": GREEN, "warn": AMBER, "danger": RED}
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fills[tone])
    set_cell_border(cell, accents[tone], 8)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(accents[tone])
    p2 = cell.add_paragraph(body, style="Body Compact")
    p2.paragraph_format.space_after = Pt(0)
    width = section_content_width_dxa(doc.sections[-1])
    apply_table_geometry(table, [width], table_width_dxa=width)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], weights: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header = table.rows[0]
    set_repeat_table_header(header)
    prevent_row_split(header)
    for idx, label in enumerate(headers):
        cell = header.cells[idx]
        set_cell_shading(cell, NAVY)
        set_cell_border(cell, NAVY, 6)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(WHITE)
        r.font.size = Pt(8.5)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        cells = row.cells
        for col_idx, value in enumerate(values):
            cell = cells[col_idx]
            set_cell_shading(cell, WHITE if row_idx % 2 == 0 else PALE_BLUE)
            set_cell_border(cell, LINE, 5)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cell.paragraphs[0]
            p.style = doc.styles["Table Text"]
            p.add_run(value)
    total = section_content_width_dxa(doc.sections[-1])
    widths = column_widths_from_weights(weights, total)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=total,
        cell_margins_dxa={"top": 50, "bottom": 50, "start": 120, "end": 120},
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def set_doc_defaults(doc: Document) -> tuple[int, int]:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.65)
    section.bottom_margin = Cm(1.55)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.75)
    section.footer_distance = Cm(0.65)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.13

    for name, size, color, before, after in [
        ("Title", 30, NAVY, 0, 10),
        ("Subtitle", 14, MUTED, 0, 8),
        ("Heading 1", 18, NAVY, 12, 5),
        ("Heading 2", 12.5, BLUE, 8, 3),
        ("Heading 3", 10.5, GREEN, 6, 2),
    ]:
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in ("Title", "Heading 1") else "Aptos"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = name != "Subtitle"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    body_compact = doc.styles.add_style("Body Compact", WD_STYLE_TYPE.PARAGRAPH)
    body_compact.base_style = normal
    body_compact.font.name = "Aptos"
    body_compact.font.size = Pt(9.2)
    body_compact.paragraph_format.space_after = Pt(3)
    body_compact.paragraph_format.line_spacing = 1.1

    table_text = doc.styles.add_style("Table Text", WD_STYLE_TYPE.PARAGRAPH)
    table_text.base_style = normal
    table_text.font.name = "Aptos"
    table_text.font.size = Pt(8.2)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.05

    code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    code.base_style = normal
    code.font.name = "Consolas"
    code.font.size = Pt(8)
    code.font.color.rgb = RGBColor.from_string(NAVY)
    code.paragraph_format.left_indent = Cm(0.35)
    code.paragraph_format.right_indent = Cm(0.35)
    code.paragraph_format.space_before = Pt(2)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.line_spacing = 1.0
    p_pr = code.element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "EEF3F6")
    p_pr.append(shd)

    bullet_id = add_numbering_definition(doc, "bullet")
    number_id = add_numbering_definition(doc, "decimal")
    return bullet_id, number_id


def set_header_footer(doc: Document) -> None:
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = "MEDCARESO 2.1.1  /  GUIDE D’INSTALLATION CABINET"
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        r = p.runs[0]
        r.font.name = "Aptos"
        r.font.size = Pt(7.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(MUTED)
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_after = Pt(0)
        fr = fp.add_run("Document opérationnel  •  Page ")
        fr.font.name = "Aptos"
        fr.font.size = Pt(7.5)
        fr.font.color.rgb = RGBColor.from_string(MUTED)
        add_page_field(fp)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    r = p.add_run("MEDCARESO")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor.from_string(GREEN)

    title = doc.add_paragraph(style="Title")
    title.add_run("Guide d’installation complète\nd’un cabinet")
    subtitle = doc.add_paragraph(style="Subtitle")
    subtitle.add_run("PostgreSQL local partagé  •  Multi-postes  •  Portail patient par QR code")

    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_callout(
        doc,
        "OBJECTIF",
        "Installer MedCareSO de manière reproductible, sécurisée et testée, depuis la préparation du poste serveur jusqu’à la remise au cabinet.",
        "ok",
    )

    add_table(
        doc,
        ["DOCUMENT", "VALEUR"],
        [
            ["Produit", "MedCareSO 2.1.1"],
            ["Public", "Installateur, administrateur du cabinet, support informatique"],
            ["Scénario", "Base PostgreSQL locale sur le PC serveur, partagée sur le réseau privé"],
            ["Mise à jour", "16 juillet 2026"],
        ],
        [1.2, 2.8],
    )

    doc.add_paragraph().paragraph_format.space_after = Pt(36)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Aucun mot de passe réel n’est inclus dans ce document.\nLes secrets sont remis séparément au responsable autorisé.")
    r.font.size = Pt(9)
    r.font.italic = True
    r.font.color.rgb = RGBColor.from_string(MUTED)
    doc.add_page_break()


def build() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.core_properties.title = "Guide d’installation MedCareSO pour un cabinet"
    doc.core_properties.subject = "Installation réseau PostgreSQL, clients, licence, QR et sauvegardes"
    doc.core_properties.author = "MedCareSO"
    doc.core_properties.keywords = "MedCareSO, installation, cabinet, PostgreSQL, réseau, QR, sauvegarde"
    bullet_id, number_id = set_doc_defaults(doc)
    add_cover(doc)

    add_heading(doc, "Le parcours en un coup d’œil", 1)
    add_table(
        doc,
        ["PHASE", "RÉSULTAT ATTENDU"],
        [
            ["1. Préparer", "Poste serveur choisi, IP réservée, réseau privé vérifié"],
            ["2. Créer la base", "PostgreSQL actif, cabinet_db et cabinet_app disponibles"],
            ["3. Ouvrir le LAN", "5432 accessible uniquement depuis le sous-réseau du cabinet"],
            ["4. Installer le serveur", "MedCareSO connecté localement et migrations appliquées"],
            ["5. Configurer", "Licence, comptes, cabinet et périphériques prêts"],
            ["6. Installer les clients", "Tous les postes utilisent la même base et la même version"],
            ["7. Activer le portail", "QR testé sur téléphone, file d’attente mise à jour"],
            ["8. Sauvegarder et valider", "Sauvegarde testée et procès-verbal de recette signé"],
        ],
        [1.2, 3.5],
    )
    add_callout(doc, "RÈGLE D’OR", "Le PC serveur doit rester allumé et MedCareSO doit rester ouvert pour que le portail patient local soit accessible.", "warn")

    add_heading(doc, "1. Architecture cible", 1)
    doc.add_paragraph(
        "Le cabinet possède une seule base de données PostgreSQL. Elle se trouve sur un poste serveur du cabinet et est partagée uniquement sur le réseau local. Tous les postes MedCareSO lisent et écrivent dans cette même base.",
        style="Body Compact",
    )
    add_table(
        doc,
        ["ÉLÉMENT", "CONTENU", "CONNEXION"],
        [
            ["PC serveur / principal", "PostgreSQL + MedCareSO + portail patient", "Base locale via 127.0.0.1:5432"],
            ["PC client", "MedCareSO uniquement", "Base distante via IP_LAN_SERVEUR:5432"],
            ["Téléphone patient", "Navigateur web, aucune application à installer", "Portail via http://IP_LAN_SERVEUR:4580/rdv/TOKEN"],
            ["Support de sauvegarde", "Copies chiffrées de la base", "Disque externe ou stockage hors du poste serveur"],
        ],
        [1.25, 1.7, 2.2],
    )
    add_callout(doc, "À NE PAS FAIRE", "Ne pas utiliser 127.0.0.1 sur un PC client : cette adresse désigne le PC client lui-même, pas le serveur.", "danger")

    add_heading(doc, "2. Avant de commencer", 1)
    add_heading(doc, "2.1 Choisir le bon scénario", 2)
    add_table(
        doc,
        ["NOUVEAU CABINET", "CABINET EXISTANT"],
        [
            ["Créer une base vide avec le script PostgreSQL complet.", "Ne jamais travailler sur l’unique base historique."],
            ["Laisser MedCareSO appliquer ses migrations au premier lancement.", "Copier la base, migrer sur un poste de test, puis comparer les données."],
            ["Créer les comptes et paramètres après validation.", "Faire une sauvegarde finale avant la bascule et conserver l’ancienne base 1 à 2 mois."],
        ],
        [1, 1],
    )
    add_heading(doc, "2.2 Fiche de préparation", 2)
    add_table(
        doc,
        ["INFORMATION", "À RENSEIGNER"],
        [
            ["Nom du cabinet / responsable", "____________________________________________"],
            ["Nom du PC serveur", "____________________________________________"],
            ["IPv4 réservée / sous-réseau", "____________________________________________"],
            ["Nombre de postes clients", "____________________________________________"],
            ["Chemin des sauvegardes", "____________________________________________"],
            ["Installateur / date", "____________________________________________"],
        ],
        [1.3, 2.7],
    )
    add_heading(doc, "2.3 Prérequis", 2)
    add_list(
        doc,
        [
            "Windows 10 ou Windows 11, 64 bits, avec un compte administrateur pour l’installation.",
            "Un poste serveur fiable, idéalement protégé par un onduleur, avec assez d’espace disque.",
            "Tous les équipements connectés au même réseau privé du cabinet ; éviter le Wi-Fi invité.",
            "L’installateur MedCareSO 2.1.1 et le dossier SETUP_WINDOWS contenant le script PostgreSQL.",
            "PostgreSQL installé sur le serveur. Node.js et npm ne sont pas nécessaires sur les postes clients.",
            "Les mots de passe administrateur PostgreSQL et cabinet_app conservés dans un support sécurisé séparé.",
        ],
        bullet_id,
    )

    add_heading(doc, "3. Préparer le réseau du cabinet", 1)
    add_list(
        doc,
        [
            "Donner un nom clair au serveur, par exemple MEDCARESO-SRV.",
            "Dans le routeur, réserver une IPv4 au serveur par DHCP afin qu’elle ne change pas.",
            "Sur le serveur, exécuter ipconfig et noter l’IPv4 et le masque. Exemple : 192.168.1.86 / 255.255.255.0.",
            "Vérifier que les postes et téléphones sont dans le même sous-réseau et que l’isolation des clients Wi-Fi est désactivée.",
            "Conserver cette adresse dans la fiche de remise ; elle sera utilisée par les postes clients et le QR code.",
        ],
        number_id,
        True,
    )
    add_code(doc, "ipconfig")
    add_callout(doc, "EXEMPLE", "Si le serveur est 192.168.1.86 et le réseau est /24, les clients sont généralement autorisés dans 192.168.1.0/24.", "info")

    add_heading(doc, "4. Installer et préparer PostgreSQL sur le serveur", 1)
    add_heading(doc, "4.1 Installer le moteur", 2)
    add_list(
        doc,
        [
            "Installer PostgreSQL et les outils en ligne de commande. pgAdmin peut être installé pour l’administration.",
            "Conserver le port 5432 sauf contrainte documentée.",
            "Définir un mot de passe fort pour l’administrateur postgres et le remettre séparément.",
            "Ouvrir services.msc et vérifier que le service PostgreSQL est En cours d’exécution et en démarrage automatique.",
        ],
        bullet_id,
    )
    add_heading(doc, "4.2 Créer cabinet_db et cabinet_app", 2)
    doc.add_paragraph(
        "Depuis un terminal administrateur, lancer le script fourni. Il crée ou met à jour le rôle cabinet_app, crée la base cabinet_db avec le bon propriétaire et attribue les privilèges nécessaires. Le script demande le mot de passe applicatif sans l’écrire en clair dans le fichier.",
        style="Body Compact",
    )
    add_code(doc, 'psql -U postgres -f "SETUP_WINDOWS\\SCRIPT_POSTGRESQL_COMPLET.sql"')
    add_table(
        doc,
        ["PARAMÈTRE", "VALEUR"],
        [
            ["Base", "cabinet_db"],
            ["Utilisateur applicatif", "cabinet_app"],
            ["Mot de passe", "Secret unique du cabinet — ne pas reprendre un exemple du projet"],
            ["Port", "5432 par défaut"],
        ],
        [1.2, 2.8],
    )
    add_heading(doc, "4.3 Tester localement", 2)
    add_code(doc, 'psql -h 127.0.0.1 -U cabinet_app -d cabinet_db -c "SELECT current_database(), current_user;"')
    add_callout(doc, "RÉSULTAT ATTENDU", "La commande retourne cabinet_db et cabinet_app. Si elle échoue, corriger PostgreSQL avant d’installer MedCareSO.", "ok")

    add_heading(doc, "5. Autoriser PostgreSQL sur le réseau local", 1)
    add_heading(doc, "5.1 postgresql.conf", 2)
    doc.add_paragraph("Dans le fichier de configuration de l’instance PostgreSQL, autoriser l’écoute réseau :", style="Body Compact")
    add_code(doc, "listen_addresses = '*'")
    doc.add_paragraph("Pour une politique plus stricte, utiliser l’IPv4 fixe du serveur à la place de *. Conserver le port 5432.", style="Body Compact")
    add_heading(doc, "5.2 pg_hba.conf", 2)
    doc.add_paragraph("Ajouter une règle limitée à la base, au compte applicatif et au sous-réseau réel du cabinet :", style="Body Compact")
    add_code(doc, "host    cabinet_db    cabinet_app    192.168.1.0/24    scram-sha-256")
    add_callout(doc, "SÉCURITÉ", "Ne pas ajouter une règle 0.0.0.0/0. Ne pas donner à cabinet_app l’accès à la base système postgres.", "danger")
    add_heading(doc, "5.3 Service et pare-feu", 2)
    add_list(
        doc,
        [
            "Redémarrer le service PostgreSQL après les modifications.",
            "Créer une règle entrante TCP 5432, profil Privé, source limitée au sous-réseau du cabinet.",
            "Depuis un PC client, vérifier le port avant de lancer l’installateur.",
        ],
        bullet_id,
    )
    add_code(doc, "Test-NetConnection 192.168.1.86 -Port 5432")
    add_callout(doc, "RÉSULTAT ATTENDU", "TcpTestSucceeded : True. Sinon, vérifier l’IPv4, le service, listen_addresses, pg_hba.conf et le pare-feu.", "ok")

    add_heading(doc, "6. Installer le poste serveur", 1)
    add_list(
        doc,
        [
            "Copier l’installateur depuis une source approuvée et vérifier sa version.",
            "Lancer l’installation avec les droits administrateur.",
            "Choisir PC serveur / principal.",
            "Renseigner Hôte 127.0.0.1, Port 5432, Base cabinet_db, Utilisateur cabinet_app et le mot de passe du cabinet.",
            "Laisser l’assistant tester la connexion. Le logiciel n’est lancé que si le test réussit.",
            "Au premier démarrage, attendre l’application automatique des migrations puis vérifier que l’écran de connexion apparaît.",
        ],
        number_id,
        True,
    )
    add_callout(doc, "SMART APP CONTROL", "En production, distribuer un installateur signé. Si Windows bloque un binaire non signé, vérifier sa provenance et suivre la procédure informatique approuvée ; ne pas désactiver durablement une protection de sécurité.", "warn")

    add_heading(doc, "7. Première ouverture et configuration du cabinet", 1)
    add_heading(doc, "7.1 Accès administrateur et licence", 2)
    add_list(
        doc,
        [
            "Se connecter avec le compte super-administrateur remis dans le dossier sécurisé, sans recopier le mot de passe dans ce guide.",
            "Activer la licence une fois sur la base du cabinet.",
            "Installer exactement la même version MedCareSO sur tous les postes : la licence est partagée par la base du cabinet.",
            "Changer immédiatement les mots de passe temporaires et créer un compte nominatif par utilisateur.",
        ],
        bullet_id,
    )
    add_heading(doc, "7.2 Comptes et spécialités", 2)
    add_list(
        doc,
        [
            "Créer les médecins, assistants et administrateurs nécessaires.",
            "Attribuer la spécialité correcte à chaque médecin afin d’afficher les modules adaptés.",
            "Tester chaque compte et vérifier les droits avant la mise en service.",
            "Éviter les comptes partagés : ils réduisent la traçabilité des actions.",
        ],
        bullet_id,
    )
    add_heading(doc, "7.3 Paramètres et périphériques", 2)
    add_table(
        doc,
        ["ZONE", "À CONFIGURER / TESTER"],
        [
            ["Général", "Coordonnées du cabinet, horaires, informations légales"],
            ["Documents", "En-tête, logo, modèles et aperçu d’impression"],
            ["Impression", "Imprimante standard et impression d’une page test"],
            ["Point de vente", "Imprimante thermique si utilisée"],
            ["Numérisation", "Scanner USB et import d’un document test"],
            ["SMS / Cloud", "Configurer uniquement si le contrat et les identifiants sont disponibles"],
        ],
        [1.2, 3.2],
    )

    add_heading(doc, "8. Installer chaque poste client", 1)
    add_list(
        doc,
        [
            "Vérifier que Test-NetConnection IP_SERVEUR -Port 5432 réussit.",
            "Lancer le même installateur MedCareSO 2.1.1 avec les droits administrateur.",
            "Choisir PC client.",
            "Renseigner l’IPv4 réservée du serveur, jamais 127.0.0.1.",
            "Utiliser la même base cabinet_db, le même utilisateur cabinet_app et le même mot de passe applicatif.",
            "Cliquer sur Tester, puis Sauvegarder. Se connecter avec un compte utilisateur nominatif.",
            "Créer un patient de test depuis un poste et vérifier sa présence depuis un autre, puis supprimer le test.",
        ],
        number_id,
        True,
    )
    add_callout(doc, "COHÉRENCE DES VERSIONS", "Après toute mise à jour du serveur, mettre à jour tous les postes clients avant de reprendre l’activité.", "warn")

    add_heading(doc, "9. Activer la page patient et le QR code", 1)
    add_list(
        doc,
        [
            "Sur le poste serveur, ouvrir Paramètres > RDV.",
            "Activer le portail local et conserver le port 4580, sauf conflit documenté.",
            "Pour un usage uniquement dans le cabinet, laisser l’URL publique vide.",
            "Enregistrer, afficher ou imprimer le QR code généré.",
            "Connecter un téléphone au même Wi-Fi privé, scanner le QR et vérifier l’ouverture de la page.",
            "Tester une arrivée avec rendez-vous et une arrivée sans rendez-vous.",
            "Vérifier que le ticket apparaît automatiquement dans la salle d’attente et que l’ordre est visible côté patient.",
        ],
        number_id,
        True,
    )
    add_code(doc, "http://<IP_SERVEUR>:4580/rdv/<TOKEN>")
    add_table(
        doc,
        ["CONTRÔLE", "ATTENDU"],
        [
            ["Nom du patient", "Saisie minimale conforme aux règles du cabinet"],
            ["Type d’arrivée", "Avec RDV ou sans RDV"],
            ["File d’attente", "Nouvel arrivant visible sans saisie manuelle"],
            ["Affichage patient", "Position / ordre sans exposer les données médicales des autres patients"],
            ["Disponibilité", "Accessible tant que le serveur et MedCareSO sont actifs"],
        ],
        [1.2, 3.1],
    )
    add_callout(doc, "CONFIDENTIALITÉ", "Le portail ne doit jamais afficher le diagnostic, le numéro de téléphone ou d’autres informations médicales d’un patient à un autre.", "danger")

    add_heading(doc, "10. Mettre en place les sauvegardes", 1)
    doc.add_paragraph(
        "La base du serveur est le point central du cabinet. Une sauvegarde quotidienne doit être automatique, contrôlée et copiée hors du disque du serveur.",
        style="Body Compact",
    )
    add_code(doc, 'pg_dump -h 127.0.0.1 -U cabinet_app -F c -f "D:\\MedCareSO_Backups\\cabinet_db_YYYY-MM-DD.backup" cabinet_db')
    add_table(
        doc,
        ["MESURE", "RÈGLE MINIMALE"],
        [
            ["Fréquence", "Quotidienne, après la fermeture du cabinet"],
            ["Rétention", "Plusieurs versions journalières, hebdomadaires et mensuelles selon la politique du cabinet"],
            ["Deuxième copie", "Support chiffré ou stockage externe au PC serveur"],
            ["Contrôle", "Taille et date du fichier vérifiées automatiquement"],
            ["Test de restauration", "Régulier, dans une base de test, par une personne autorisée"],
        ],
        [1.15, 3.2],
    )
    add_callout(doc, "IMPORTANT", "Une copie qui n’a jamais été restaurée n’est pas une sauvegarde prouvée. Documenter la date et le résultat de chaque test de restauration.", "warn")

    add_heading(doc, "11. Recette avant mise en service", 1)
    add_table(
        doc,
        ["✓", "TEST", "RÉSULTAT / OBSERVATION"],
        [
            ["☐", "Redémarrage serveur : PostgreSQL et MedCareSO fonctionnent", "____________________________"],
            ["☐", "Connexion locale au serveur", "____________________________"],
            ["☐", "Connexion de chaque poste client", "____________________________"],
            ["☐", "Même patient visible sur deux postes", "____________________________"],
            ["☐", "Licence reconnue sur tous les postes", "____________________________"],
            ["☐", "Comptes et permissions vérifiés", "____________________________"],
            ["☐", "Imprimante et scanner testés", "____________________________"],
            ["☐", "QR patient testé avec et sans RDV", "____________________________"],
            ["☐", "File d’attente mise à jour automatiquement", "____________________________"],
            ["☐", "Sauvegarde créée et restauration de test validée", "____________________________"],
        ],
        [0.3, 2.2, 1.8],
    )

    add_heading(doc, "12. Dépannage rapide", 1)
    add_table(
        doc,
        ["SYMPTÔME", "CAUSE PROBABLE", "ACTION"],
        [
            ["Timeout PostgreSQL", "Port, service, IP ou pare-feu", "Tester 5432, confirmer l’IP réservée, le service et la règle Privé"],
            ["no pg_hba.conf entry", "Sous-réseau ou base non autorisé", "Ajouter une règle cabinet_db / cabinet_app limitée au sous-réseau, puis redémarrer PostgreSQL"],
            ["Erreur mentionnant database postgres", "Ancienne version ou configuration obsolète", "Mettre à jour tous les postes ; ne pas ouvrir l’accès de cabinet_app à la base système"],
            ["Tester fonctionne, Sauvegarder échoue", "Un ancien composant relance un test vers une mauvaise base", "Vérifier la version sur ce PC et réinstaller la même version que le serveur"],
            ["Nom ou mot de passe incorrect", "Poste connecté à une autre base ou compte absent", "Contrôler le fichier de configuration, la base cabinet_db et la création du compte"],
            ["Licence retirée d’un autre PC", "Versions incompatibles / ancien système de licence", "Installer la même version actuelle sur tous les postes et réactiver une fois"],
            ["QR inaccessible", "Application fermée, mauvais Wi-Fi, IP changée ou port 4580 bloqué", "Ouvrir MedCareSO sur le serveur, vérifier IP, Wi-Fi privé et pare-feu"],
            ["Application se ferme immédiatement", "ELECTRON_RUN_AS_NODE=1 dans l’environnement Windows", "Supprimer la variable utilisateur, puis fermer et relancer l’application"],
            ["Windows bloque l’installateur", "Binaire non signé ou non reconnu", "Vérifier la source et utiliser un installateur signé / procédure approuvée"],
        ],
        [1.3, 1.45, 2.15],
    )
    add_heading(doc, "Commande de correction Electron", 2)
    add_code(doc, "$env:ELECTRON_RUN_AS_NODE\n[Environment]::SetEnvironmentVariable('ELECTRON_RUN_AS_NODE', $null, 'User')")

    add_heading(doc, "13. Cas d’un cabinet existant : migration contrôlée", 1)
    add_list(
        doc,
        [
            "Faire une copie complète de l’ancienne base avant toute opération.",
            "Installer et migrer d’abord sur un poste de test, jamais directement sur l’unique production.",
            "Comparer les volumes : patients, rendez-vous, plans de traitement, paiements, dettes et inventaire.",
            "Ouvrir manuellement trois ou quatre dossiers patients connus et comparer les informations clés.",
            "Faire une sauvegarde finale juste avant la bascule, puis arrêter les écritures dans l’ancien système.",
            "Conserver l’ancienne base en lecture seule pendant un à deux mois selon la politique du cabinet.",
        ],
        number_id,
        True,
    )
    add_callout(doc, "GO / NO-GO", "Ne mettre le nouveau système en production que lorsque les contrôles de données, la sauvegarde et la restauration sont tous validés.", "danger")

    add_heading(doc, "14. Remise au cabinet", 1)
    add_table(
        doc,
        ["ÉLÉMENT REMIS", "CONTENU"],
        [
            ["Fiche technique", "Nom et IPv4 du serveur, port 5432, base cabinet_db, version MedCareSO"],
            ["Secrets", "Mots de passe remis séparément au responsable autorisé"],
            ["Installation", "Copie vérifiée de l’installateur et, si disponible, son empreinte SHA-256"],
            ["Sauvegarde", "Emplacement, fréquence, rétention, responsable et dernière restauration testée"],
            ["Portail", "QR imprimé, URL locale et consigne de maintenir le serveur actif"],
            ["Support", "Contact, horaires et procédure en cas d’incident"],
        ],
        [1.15, 3.25],
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    add_table(
        doc,
        ["VALIDATION", "NOM / DATE / SIGNATURE"],
        [
            ["Installateur", "________________________________________________"],
            ["Responsable du cabinet", "________________________________________________"],
        ],
        [1.2, 2.8],
    )
    add_callout(doc, "INSTALLATION TERMINÉE", "Le cabinet est prêt lorsque les tests de la section 11 sont validés, la sauvegarde est opérationnelle et le responsable a reçu les informations de remise.", "ok")

    set_header_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    result = build()
    print(result)
